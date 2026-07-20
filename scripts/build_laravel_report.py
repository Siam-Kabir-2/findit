"""Build a simple Laravel-focused FindIt project report with current screenshots."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(r"d:\3-1\db+web\findit")
SHOTS = ROOT / "docs" / "presentation" / "screenshots"
OUT_DOWNLOADS = Path(r"c:\Users\siams\Downloads\FindIt_Project_Report.docx")
OUT_PROJECT = ROOT / "docs" / "FindIt_Project_Report.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_title(doc, text, size=22):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=RGBColor(0x1E, 0x29, 0x3B))
    return p


def add_subtitle(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=12, color=RGBColor(0x47, 0x55, 0x69))
    return p


def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        set_run_font(run, size=14, bold=True, color=RGBColor(0x0F, 0x17, 0x2A))
    return p


def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        set_run_font(run, size=12, bold=True, color=RGBColor(0x1E, 0x29, 0x3B))
    return p


def add_p(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=11)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run, size=11)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = val
            for p in cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, size=10)
    doc.add_paragraph()


def add_figure(doc, filename, caption, width=5.8):
    path = SHOTS / filename
    if not path.exists():
        add_p(doc, f"[Missing screenshot: {filename}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9, color=RGBColor(0x64, 0x74, 0x8B))
    cap.paragraph_format.space_after = Pt(12)


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    add_title(doc, "FindIt Project Report")
    add_subtitle(doc, "Campus Lost & Found Web Application")
    add_subtitle(doc, "Laravel 12 · Blade · MySQL · Vite")
    doc.add_paragraph()

    # 1
    add_h1(doc, "1. Introduction")
    add_p(
        doc,
        "FindIt is a campus lost-and-found website. Students can report lost or found items, "
        "browse listings, and submit ownership claims. Administrators review claims, manage "
        "catalog data, and monitor activity from a separate admin console.",
    )
    add_p(
        doc,
        "This report focuses on website development with Laravel: routes, Blade views, dual "
        "authentication, forms, image uploads, campus map pins, and how the app uses MySQL. "
        "Database rules support the site, but the main product is the web application.",
    )

    # 2
    add_h1(doc, "2. Project Objectives")
    add_bullets(
        doc,
        [
            "Build a clean Laravel web app for campus lost & found workflows.",
            "Provide public browse/detail pages and authenticated user features.",
            "Implement a separate admin console with its own login guard.",
            "Support item photos, KUET location selection, and map pins on item pages.",
            "Keep business rules in a PHP service layer with MySQL audit triggers.",
            "Ship a simple, modern UI with Blade, Vite, and Tailwind CSS.",
        ],
    )

    # 3
    add_h1(doc, "3. System Architecture")
    add_p(
        doc,
        "The browser renders Blade pages. Laravel handles routing, validation, sessions, and "
        "file uploads. Business writes go through FinditPlsqlService (PHP transactions). "
        "MySQL stores data; triggers write audit rows when items or claims change.",
    )
    add_table(
        doc,
        ["Layer", "Role"],
        [
            ["Laravel 12 + Blade", "Pages, forms, auth, controllers, validation"],
            ["FinditPlsqlService", "Item/claim business rules and transactions"],
            ["MySQL (XAMPP)", "Persistent data + audit triggers"],
            ["Vite + Tailwind", "Frontend assets and styling"],
            ["Leaflet + OSM", "Campus map pins for KUET locations"],
        ],
    )
    add_figure(doc, "26-architecture.png", "Figure 1. High-level architecture", width=5.4)

    # 4
    add_h1(doc, "4. Website Features")
    add_h2(doc, "4.1 Public site")
    add_bullets(
        doc,
        [
            "Landing page with brand hero, stats, how-it-works, and recent items.",
            "Browse board with search and filters (type, status, category, location).",
            "Item detail with photo, metadata, KUET map pin, and claim entry.",
            "User login and registration pages.",
        ],
    )
    add_h2(doc, "4.2 Authenticated users")
    add_bullets(
        doc,
        [
            "Personal dashboard overview.",
            "Report lost/found items with optional image upload.",
            "Manage own listings; submit and track claims.",
        ],
    )
    add_h2(doc, "4.3 Admin console")
    add_bullets(
        doc,
        [
            "Separate admin login and session.",
            "Dashboard with live stats and recent activity.",
            "Approve/reject claims; manage items, users, categories, locations.",
            "Audit timeline generated from database triggers.",
        ],
    )

    # 5
    add_h1(doc, "5. Laravel Implementation (Brief)")
    add_h2(doc, "5.1 Structure")
    add_p(
        doc,
        "Controllers live under app/Http/Controllers (user + Admin). Models use Eloquent. "
        "Views are Blade templates under resources/views. Routes are declared in routes/web.php. "
        "Frontend assets compile through Vite from resources/css and resources/js.",
    )
    add_h2(doc, "5.2 Authentication")
    add_p(
        doc,
        "Two guards are used: web for students and admin for administrators. This keeps "
        "student sessions and admin sessions separate and avoids shared login confusion.",
    )
    add_h2(doc, "5.3 Key workflows")
    add_bullets(
        doc,
        [
            "Report item → validate form → store image → create item via service.",
            "Submit claim → block self-claims / closed items → insert PENDING claim.",
            "Approve claim → approve one claim, reject other pending claims, mark item CLAIMED.",
            "Locations → KUET campus list with coordinates → map shown on item detail.",
        ],
    )
    add_figure(doc, "29-workflow.png", "Figure 2. Main application workflow", width=5.4)

    # 6 Data layer short
    add_h1(doc, "6. Data Layer Supporting the Website")
    add_p(
        doc,
        "The site runs on MySQL. Core tables are users, admins, categories, locations, items, "
        "claims, and audit_logs. Locations include latitude/longitude for KUET map pins. "
        "AFTER INSERT/UPDATE/DELETE triggers on items and claims write audit_logs so admins "
        "can see what changed without extra application code.",
    )
    add_table(
        doc,
        ["Table", "Used by the website for"],
        [
            ["users / admins", "Login and ownership"],
            ["categories / locations", "Filters, forms, map pins"],
            ["items", "Browse board and detail pages"],
            ["claims", "User claims and admin review"],
            ["audit_logs", "Admin activity timeline"],
        ],
    )

    # 7 Screenshots - main body
    add_h1(doc, "7. Website Screenshots")
    add_p(doc, "Updated screenshots from the running Laravel application.")

    add_h2(doc, "7.1 Public pages")
    add_figure(doc, "20-home-viewport.png", "Figure 3. Home page (first viewport)", width=5.9)
    add_figure(doc, "01-home.png", "Figure 4. Home page (full)", width=5.6)
    add_figure(doc, "19-browse-viewport.png", "Figure 5. Browse board", width=5.9)
    add_figure(doc, "02-browse-items.png", "Figure 6. Browse with filters and listings", width=5.6)
    add_figure(doc, "03-item-detail.png", "Figure 7. Item detail with KUET map", width=5.6)
    add_figure(doc, "04-user-login.png", "Figure 8. User login", width=5.4)
    add_figure(doc, "05-register.png", "Figure 9. User registration", width=5.4)

    add_h2(doc, "7.2 User area")
    add_figure(doc, "06-user-dashboard.png", "Figure 10. User dashboard", width=5.6)
    add_figure(doc, "07-report-item.png", "Figure 11. Report item form (KUET locations)", width=5.4)
    add_figure(doc, "08-my-items.png", "Figure 12. My items", width=5.4)
    add_figure(doc, "09-my-claims.png", "Figure 13. My claims", width=5.4)
    add_figure(doc, "23-claim-form.png", "Figure 14. Claim form on item detail", width=5.4)
    add_figure(doc, "24-claim-form-filled.png", "Figure 15. Claim form filled", width=5.4)

    add_h2(doc, "7.3 Admin console")
    add_figure(doc, "10-admin-login.png", "Figure 16. Admin login", width=5.4)
    add_figure(doc, "11-admin-dashboard.png", "Figure 17. Admin dashboard", width=5.6)
    add_figure(doc, "12-admin-claims.png", "Figure 18. Claims management", width=5.6)
    add_figure(doc, "14-admin-items.png", "Figure 19. Items management", width=5.6)
    add_figure(doc, "15-admin-users.png", "Figure 20. Users management", width=5.6)
    add_figure(doc, "16-admin-categories.png", "Figure 21. Categories", width=5.2)
    add_figure(doc, "17-admin-locations.png", "Figure 22. KUET locations + map pins", width=5.6)
    add_figure(doc, "18-admin-audit.png", "Figure 23. Audit activity timeline", width=5.6)

    # 8 Testing
    add_h1(doc, "8. Testing Checklist")
    add_table(
        doc,
        ["Area", "What was checked"],
        [
            ["Public UI", "Home, browse filters, item detail, map pin visible"],
            ["Auth", "User register/login; admin login on separate guard"],
            ["User flows", "Report item, my items, submit claim"],
            ["Admin flows", "Approve/reject claim, manage catalog, open audit"],
            ["Assets", "Vite build loads CSS/JS; uploaded images display"],
        ],
    )

    # 9 Conclusion
    add_h1(doc, "9. Conclusion")
    add_p(
        doc,
        "FindIt is a practical Laravel website for campus lost-and-found: public discovery, "
        "authenticated reporting/claiming, and an admin console for verification. MySQL and "
        "triggers support the app with reliable storage and an audit trail, while the user "
        "experience is delivered through Blade pages, dual auth, forms, images, and KUET maps.",
    )

    OUT_DOWNLOADS.parent.mkdir(parents=True, exist_ok=True)
    OUT_PROJECT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOWNLOADS)
    doc.save(OUT_PROJECT)
    print(f"Wrote {OUT_DOWNLOADS}")
    print(f"Wrote {OUT_PROJECT}")


if __name__ == "__main__":
    build()

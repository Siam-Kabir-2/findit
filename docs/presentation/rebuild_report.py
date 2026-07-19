"""Rebuild FindIt project report: DB-focused, with screenshots."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement

ROOT = Path(r"d:\3-1\db+web\findit")
SHOTS = ROOT / "docs" / "presentation" / "screenshots"
OUT = Path(r"c:\Users\siams\Downloads\FindIt_Project_Report.docx")
BACKUP = Path(r"c:\Users\siams\Downloads\FindIt_Project_Report_backup.docx")


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    return p


def add_para(doc, text, *, bold=False, size=11, space_after=8):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(2)


def set_cell_shading(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        set_run_font(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_shading(hdr[i], "1A3A5C")
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_run_font(run, size=10)
            if r_i % 2 == 1:
                set_cell_shading(cell, "F0F4F8")
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def add_image(doc, name, caption, width=5.8):
    path = SHOTS / name
    if not path.exists():
        add_para(doc, f"[Missing image: {name}]", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    set_run_font(r, size=9, bold=False, color=RGBColor(0x55, 0x55, 0x55))
    cap.paragraph_format.space_after = Pt(14)


def build():
    if OUT.exists() and not BACKUP.exists():
        shutil.copy2(OUT, BACKUP)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("FindIt Project Report")
    set_run_font(r, size=22, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run("Campus Lost & Found System — Oracle Database Focus")
    set_run_font(r, size=13, color=RGBColor(0x44, 0x44, 0x44))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Database Management Systems Course Project")
    set_run_font(r, size=11, color=RGBColor(0x66, 0x66, 0x66))
    doc.add_paragraph()

    # 1. Introduction
    add_heading(doc, "1. Introduction", 1)
    add_para(
        doc,
        "FindIt is a campus lost-and-found management system. Students report lost or found "
        "items, submit ownership claims, and administrators review those claims. The web "
        "interface is built with Laravel, but the core of this project is the Oracle Database "
        "11g XE design: normalized tables, integrity constraints, sequences, triggers, "
        "audit logging, and a PL/SQL package that enforces business rules.",
    )
    add_para(
        doc,
        "This report focuses on the database layer — schema, relationships, PL/SQL, and how "
        "database rules drive the application — with screenshots of the working system.",
    )

    # 2. Objectives
    add_heading(doc, "2. Project Objectives", 1)
    add_bullets(
        doc,
        [
            "Design a normalized Oracle schema for users, items, claims, categories, locations, and audit history.",
            "Enforce data integrity with primary keys, foreign keys, UNIQUE, and CHECK constraints.",
            "Generate surrogate keys using sequences and BEFORE INSERT triggers.",
            "Centralize business logic in a PL/SQL package (findit_pkg).",
            "Record item and claim changes automatically with AFTER triggers into audit_logs.",
            "Demonstrate atomic multi-row workflows (especially claim approval) inside the database.",
        ],
    )

    # 3. Architecture (brief)
    add_heading(doc, "3. System Architecture (Brief)", 1)
    add_para(
        doc,
        "The browser uses Blade pages. Laravel validates input and calls Oracle. Critical "
        "writes go through findit_pkg procedures; reads use SQL/joins. Sessions stay in the "
        "application layer so the database remains focused on data and business rules.",
    )
    add_image(doc, "26-architecture.png", "Figure 1. Layered architecture — UI → Laravel → Oracle / findit_pkg")
    add_image(doc, "29-workflow.png", "Figure 2. Main workflow: report → claim → admin decision → audit")

    add_table(
        doc,
        ["Layer", "Role"],
        [
            ["Oracle 11g XE", "Tables, constraints, sequences, triggers, findit_pkg"],
            ["PL/SQL package", "Register users, add items, submit/approve/reject claims, stats"],
            ["Laravel + PDO_OCI", "Forms, auth guards, calls procedures, displays query results"],
        ],
        [1.8, 4.5],
    )

    # 4. Database design
    add_heading(doc, "4. Database Design", 1)
    add_para(
        doc,
        "The schema has seven tables. Categories and locations are lookup tables (3NF style). "
        "Claims are a separate entity so many users can claim one item. Admins are stored "
        "separately from normal users. Audit logs keep history without foreign keys so records "
        "survive deletes.",
    )

    add_heading(doc, "4.1 Entity-Relationship Diagram", 2)
    add_image(doc, "27-er-diagram.png", "Figure 3. ER diagram of the FindIt Oracle schema")

    add_heading(doc, "4.2 Relational Schema", 2)
    add_image(doc, "28-schema-diagram.png", "Figure 4. Relational schema with keys and relationships")

    add_heading(doc, "4.3 Tables Overview", 2)
    add_image(doc, "33-oracle-tables.png", "Figure 5. Oracle objects — tables created for FindIt")

    add_table(
        doc,
        ["Table", "Purpose", "Key constraints"],
        [
            ["USERS", "Campus user accounts", "PK user_id; UNIQUE email; CHECK email LIKE '%@%'"],
            ["ADMINS", "Administrator accounts", "PK admin_id; UNIQUE email"],
            ["CATEGORIES", "Item category lookup", "PK category_id; UNIQUE name"],
            ["LOCATIONS", "Campus location lookup", "PK location_id; UNIQUE name"],
            ["ITEMS", "Lost/found reports", "FKs to users, categories, locations; CHECK type/status"],
            ["CLAIMS", "Ownership claims", "FKs to items, users; CHECK claim_status"],
            ["AUDIT_LOGS", "Change history", "CHECK action_type; no FK (history survives deletes)"],
        ],
        [1.4, 2.0, 3.0],
    )

    add_heading(doc, "4.4 Important Columns", 2)
    add_para(doc, "ITEMS", bold=True, space_after=4)
    add_table(
        doc,
        ["Column", "Notes"],
        [
            ["item_id", "Primary key (sequence + trigger)"],
            ["user_id / category_id / location_id", "Foreign keys (ON DELETE CASCADE from parents)"],
            ["item_type", "CHECK: LOST or FOUND"],
            ["status", "CHECK: PENDING, FOUND, CLAIMED, RETURNED, REJECTED"],
            ["item_image, lost_or_found_date", "Optional image path and event date"],
        ],
        [2.4, 4.0],
    )

    add_para(doc, "CLAIMS", bold=True, space_after=4)
    add_table(
        doc,
        ["Column", "Notes"],
        [
            ["claim_id", "Primary key (sequence + trigger)"],
            ["item_id, user_id", "Foreign keys to the item and claimant"],
            ["claim_message, proof_description", "Claim text and ownership proof"],
            ["claim_status", "CHECK: PENDING, APPROVED, REJECTED"],
        ],
        [2.4, 4.0],
    )

    add_para(doc, "AUDIT_LOGS", bold=True, space_after=4)
    add_table(
        doc,
        ["Column", "Notes"],
        [
            ["table_name, record_id", "Which table/row changed (soft reference)"],
            ["action_type", "INSERT, UPDATE, or DELETE"],
            ["old_status, new_status", "Status transition for items/claims"],
            ["action_by, action_time", "Who/when (package may store admin name)"],
        ],
        [2.4, 4.0],
    )

    add_heading(doc, "4.5 Relationships", 2)
    add_table(
        doc,
        ["Parent", "Cardinality", "Child"],
        [
            ["USERS", "1 : N", "ITEMS (reporter)"],
            ["USERS", "1 : N", "CLAIMS (claimant)"],
            ["CATEGORIES", "1 : N", "ITEMS"],
            ["LOCATIONS", "1 : N", "ITEMS"],
            ["ITEMS", "1 : N", "CLAIMS"],
        ],
        [1.8, 1.5, 3.0],
    )
    add_para(
        doc,
        "Deleting a user cascades to their items and claims. Deleting a category or location "
        "is blocked in PL/SQL if items still reference it — safer than cascading away catalog data.",
    )

    # 5. Sequences & triggers
    add_heading(doc, "5. Sequences and Triggers", 1)
    add_para(
        doc,
        "Each table has a sequence (seq_users, seq_items, …). BEFORE INSERT triggers assign "
        "the next value when the primary key is null, and set default status PENDING for items "
        "and claims.",
    )
    add_para(
        doc,
        "AFTER INSERT/UPDATE/DELETE triggers on ITEMS and CLAIMS write rows into AUDIT_LOGS. "
        "This gives administrators a history of status changes without relying on the application "
        "to insert audit rows manually.",
    )
    add_image(doc, "18-admin-audit.png", "Figure 6. Admin audit log — rows produced by database triggers")

    # 6. PL/SQL
    add_heading(doc, "6. PL/SQL Package: findit_pkg", 1)
    add_para(
        doc,
        "Business writes are encapsulated in package findit_pkg. Laravel calls procedures through "
        "PDO anonymous blocks (BEGIN findit_pkg....; END;). Putting rules in the database means "
        "they apply even if the client is wrong or bypassed.",
    )
    add_image(doc, "34-plsql-package.png", "Figure 7. findit_pkg source — procedures and functions")
    add_image(doc, "32-oracle-sql-files.png", "Figure 8. Oracle SQL scripts in the repository (run in order)")

    add_heading(doc, "6.1 Procedures and Functions", 2)
    add_table(
        doc,
        ["Name", "Type", "Purpose"],
        [
            ["register_user", "Procedure", "Create user; reject duplicate email"],
            ["add_item", "Procedure", "Insert lost/found item; validate type"],
            ["submit_claim", "Procedure", "Create claim with ownership checks"],
            ["approve_claim", "Procedure", "Approve one claim; reject others; set item CLAIMED"],
            ["reject_claim", "Procedure", "Reject a claim; write admin audit note"],
            ["update_item_status", "Procedure", "Change item status within allowed values"],
            ["add/delete_category|location", "Procedure", "Catalog maintenance with safe deletes"],
            ["get_total_users / items / …", "Function", "Dashboard counts for admin stats"],
        ],
        [2.2, 1.1, 3.2],
    )
    add_image(doc, "11-admin-dashboard.png", "Figure 9. Admin dashboard — counts from PL/SQL functions")

    add_heading(doc, "6.2 Claim Approval (Transaction Example)", 2)
    add_para(
        doc,
        "approve_claim is the best example of an atomic database transaction. In one procedure call it:",
    )
    add_bullets(
        doc,
        [
            "Marks the selected claim as APPROVED.",
            "Rejects other PENDING claims on the same item.",
            "Sets the item status to CLAIMED.",
            "Writes an extra audit entry with the administrator name.",
        ],
    )
    add_para(
        doc,
        "All of that succeeds or fails together inside Oracle — one DB round-trip from the app.",
    )
    add_image(doc, "25-admin-claims-pending.png", "Figure 10. Admin claim review — approve/reject via findit_pkg")

    add_heading(doc, "6.3 Business Rules Enforced in PL/SQL", 2)
    add_bullets(
        doc,
        [
            "Email must be unique (register_user raises application error if duplicate).",
            "Item type must be LOST or FOUND; status must be from the allowed set.",
            "A user cannot claim their own item.",
            "Only one PENDING claim per user per item.",
            "Items already CLAIMED, RETURNED, or REJECTED cannot receive new claims.",
            "Category/location delete refused while referenced by items.",
        ],
    )

    # 7. Integrity summary
    add_heading(doc, "7. Integrity Constraints Summary", 1)
    add_table(
        doc,
        ["Mechanism", "Example in FindIt"],
        [
            ["Entity integrity", "Primary keys on all tables via sequences"],
            ["Referential integrity", "FKs from ITEMS and CLAIMS to parent tables"],
            ["Domain integrity", "CHECK on item_type, status, claim_status, email format"],
            ["Key integrity", "UNIQUE on emails, category_name, location_name"],
            ["Business integrity", "findit_pkg rules + RAISE_APPLICATION_ERROR"],
            ["Auditability", "Triggers populate AUDIT_LOGS on item/claim changes"],
        ],
        [2.0, 4.5],
    )

    # 8. Scripts
    add_heading(doc, "8. Oracle Scripts", 1)
    add_para(doc, "Scripts under database/oracle/ are run in order against Oracle 11g XE:")
    add_table(
        doc,
        ["Script", "Purpose"],
        [
            ["01_create_user_schema.sql", "Create findit user and grants"],
            ["02_create_tables.sql", "Tables, PKs, FKs, CHECKs, sequences"],
            ["03_insert_sample_data.sql", "Demo users, items, claims"],
            ["04_basic_queries.sql", "JOIN / GROUP BY demonstration queries"],
            ["05_plsql_triggers_package.sql", "Triggers + findit_pkg package"],
        ],
        [2.6, 4.0],
    )

    # 9. Application screenshots (short)
    add_heading(doc, "9. Application Screenshots (Supporting UI)", 1)
    add_para(
        doc,
        "The Laravel UI exists to demonstrate the database end-to-end. Below are a few screens "
        "only — full page-by-page web documentation is intentionally omitted.",
    )
    add_image(doc, "01-home.png", "Figure 11. Home page — recent items from Oracle", width=5.5)
    add_image(doc, "02-browse-items.png", "Figure 12. Browse/filter items (SQL joins + filters)", width=5.5)
    add_image(doc, "03-item-detail.png", "Figure 13. Item detail before submitting a claim", width=5.5)
    add_image(doc, "24-claim-form-filled.png", "Figure 14. Claim form — data sent to submit_claim", width=5.5)
    add_image(doc, "07-report-item.png", "Figure 15. Report item — data sent to add_item", width=5.5)
    add_image(doc, "06-user-dashboard.png", "Figure 16. User dashboard — personal items and claims", width=5.5)

    # 10. Testing
    add_heading(doc, "10. Database-Oriented Testing", 1)
    add_table(
        doc,
        ["Test", "Expected database result"],
        [
            ["Register duplicate email", "ORA error / package rejects; no new USERS row"],
            ["Add LOST item", "ITEMS row + INSERT audit; status PENDING"],
            ["Claim own item", "submit_claim raises error"],
            ["Second pending claim same user/item", "Rejected by package rule"],
            ["Approve claim", "One APPROVED, siblings REJECTED, item CLAIMED, audits written"],
            ["Delete category with items", "delete_category fails while references exist"],
            ["Check user_objects", "FINDIT_PKG and triggers show STATUS = VALID"],
        ],
        [2.5, 4.0],
    )

    # 11. Conclusion
    add_heading(doc, "11. Conclusion", 1)
    add_para(
        doc,
        "FindIt demonstrates a complete Oracle-backed lost-and-found design for a DBMS course: "
        "a normalized seven-table schema, integrity constraints, sequence-driven keys, "
        "trigger-based auditing, and a PL/SQL package that owns critical business workflows. "
        "The web application is a thin client that proves those database features work in practice.",
    )

    # References
    add_heading(doc, "12. References", 1)
    add_bullets(
        doc,
        [
            "FindIt repository: https://github.com/Siam-Kabir-2/findit",
            "Oracle Database / PL/SQL documentation: https://docs.oracle.com",
            "Laravel documentation (application layer only): https://laravel.com/docs",
        ],
    )

    doc.save(OUT)
    print(f"Saved: {OUT}")
    if BACKUP.exists():
        print(f"Backup: {BACKUP}")


if __name__ == "__main__":
    build()
